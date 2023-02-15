#Een pytyon file om een txt tekst om te zetten in een xlsx tekst

#packages
import docx #python-docx
import re
import pandas

document = docx.Document('Biocheck Breeders_French.docx')

Title = []
Helptext = []
Description = []
ConditionalHelpText = []
Answer1 = []
Answer2 = []
Answer3 = []
Answer4 = []
Answer5 = []
Answer6 = []
Answer7 = []
EmptyColumn = []

offset = 0
teller = 0
passeer = False

for i in range(len(document.paragraphs)):
    if re.search('Belgium', document.paragraphs[i].text):
        i += 1
        passeer = True
        while document.paragraphs[i].text == '':
            offset += 1
            i += 1
    if passeer == True:
        i = i + offset
        titel = document.paragraphs[i].text
        ret = re.compile(r'^[A-Z0-9] \. (.*)') #Hier zit ik vast!!!!
        titletext = re.findall(ret, titel)
        titletext[0] = titletext[0].replace(u'\xa0', u' ')
        Title.append(titletext)
        teller += 1
        if teller > 10:
            break

Title.pop(0)
print(Title)

questionnaire = 'Breeders' #survey name
taal = 'fr'
helplinkoud = '/en/'
helplinknieuw = '/fr/'

#titels = vragen.Title.values
#helpteksten = vragen.HelpText.values
#descriptions = vragen.Description.values
#Condtekst = vragen.ConditionalHelpText.values